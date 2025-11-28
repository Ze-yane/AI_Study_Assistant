import streamlit as st
import io
import re
import os
from typing import List, Tuple, Optional

# Optional dependencies
try:
    from fpdf import FPDF
except Exception:
    FPDF = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from PyPDF2 import PdfReader
except Exception:
    PdfReader = None


class CVBuilder:
    """A robust CV builder with templates, sections, preview and export.

    Features:
    - Fill form fields for common resume sections (personal, summary, skills, experience, education, projects)
    - Choose a template (Simple, Modern, ATS-friendly)
    - Upload a TXT/PDF resume to prefill fields (best-effort parsing)
    - Download as PDF or DOCX with safe fallbacks
    """

    def __init__(self):
        # session state keys for dynamic lists
        if "cv_experiences" not in st.session_state:
            st.session_state["cv_experiences"] = [""]
        if "cv_educations" not in st.session_state:
            st.session_state["cv_educations"] = [""]
        if "cv_projects" not in st.session_state:
            st.session_state["cv_projects"] = [""]

    # ---------------------- Utilities ----------------------
    @staticmethod
    def _safe_extract_pdf_text(uploaded_file) -> str:
        if PdfReader is None:
            return ""
        try:
            uploaded_file.seek(0)
        except Exception:
            pass
        try:
            reader = PdfReader(uploaded_file)
            text = []
            for p in reader.pages:
                t = p.extract_text() or ""
                text.append(t)
            return "\n".join(text).strip()
        except Exception:
            return ""

    @staticmethod
    def _extract_txt(uploaded_file) -> str:
        try:
            uploaded_file.seek(0)
        except Exception:
            pass
        try:
            return uploaded_file.read().decode("utf-8") if hasattr(uploaded_file, "read") else str(uploaded_file)
        except Exception:
            return ""

    @staticmethod
    def _find_email(text: str) -> Optional[str]:
        m = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
        return m.group(0) if m else None

    @staticmethod
    def _find_phone(text: str) -> Optional[str]:
        m = re.search(r"[+]?\d[\d\s().-]{5,}\d", text)
        return m.group(0) if m else None

    @staticmethod
    def _validate_email(email: str) -> bool:
        if not email:
            return False
        pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        return re.match(pattern, email) is not None

    # ---------------------- Exporters ----------------------
    def _build_pdf_bytes(self, data: dict, template: str = "Simple") -> Optional[bytes]:
        if FPDF is None:
            return None

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=12)
        pdf.add_page()

        # Header
        pdf.set_font("Arial", 'B', 18)
        pdf.cell(0, 8, data.get("name", ""), ln=True)
        pdf.set_font("Arial", size=10)
        contact = f"{data.get('email','')} | {data.get('phone','')} | {data.get('location','')}".strip(' |')
        pdf.cell(0, 6, contact, ln=True)
        pdf.ln(6)

        # Summary
        if data.get("summary"):
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 6, "Professional Summary", ln=True)
            pdf.set_font("Arial", size=10)
            pdf.multi_cell(0, 5, data.get("summary", ""))
            pdf.ln(3)

        # Skills (inline)
        skills = data.get("skills", [])
        if skills:
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 6, "Skills", ln=True)
            pdf.set_font("Arial", size=10)
            pdf.multi_cell(0, 5, ", ".join(skills))
            pdf.ln(3)

        # Experiences
        exps = data.get("experiences", [])
        if exps:
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 6, "Work Experience", ln=True)
            pdf.set_font("Arial", size=10)
            for e in exps:
                pdf.multi_cell(0, 5, f"• {e}")
            pdf.ln(3)

        # Education
        eds = data.get("educations", [])
        if eds:
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 6, "Education", ln=True)
            pdf.set_font("Arial", size=10)
            for ed in eds:
                pdf.multi_cell(0, 5, f"• {ed}")
            pdf.ln(3)

        # Projects
        prjs = data.get("projects", [])
        if prjs:
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 6, "Projects", ln=True)
            pdf.set_font("Arial", size=10)
            for p in prjs:
                pdf.multi_cell(0, 5, f"• {p}")
            pdf.ln(3)

        return pdf.output(dest='S').encode('latin-1')

    def _build_docx_bytes(self, data: dict, template: str = "Simple") -> Optional[bytes]:
        if Document is None:
            return None

        doc = Document()
        # Header
        doc.add_heading(data.get("name", ""), level=0)
        contact = f"{data.get('email','')} | {data.get('phone','')} | {data.get('location','')}".strip(' |')
        doc.add_paragraph(contact)

        # Summary
        if data.get("summary"):
            doc.add_heading('Professional Summary', level=1)
            doc.add_paragraph(data.get('summary', ''))

        # Skills
        if data.get('skills'):
            doc.add_heading('Skills', level=1)
            doc.add_paragraph(', '.join(data.get('skills')))

        # Experience
        if data.get('experiences'):
            doc.add_heading('Work Experience', level=1)
            for e in data.get('experiences'):
                doc.add_paragraph(e, style='List Bullet')

        # Education
        if data.get('educations'):
            doc.add_heading('Education', level=1)
            for ed in data.get('educations'):
                doc.add_paragraph(ed, style='List Bullet')

        # Projects
        if data.get('projects'):
            doc.add_heading('Projects', level=1)
            for p in data.get('projects'):
                doc.add_paragraph(p, style='List Bullet')

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    # ---------------------- UI ----------------------
    def ui(self):
        st.subheader("📄 CV Builder - Professional Resume Maker")

        col_left, col_right = st.columns([2, 1])

        with col_left:
            # Upload to prefill
            st.markdown("**Upload an existing resume (optional)**")
            uploaded = st.file_uploader("Upload PDF or TXT to prefill fields", type=["pdf", "txt"])
            pref_text = ""
            if uploaded:
                if uploaded.type == "application/pdf":
                    pref_text = self._safe_extract_pdf_text(uploaded)
                else:
                    pref_text = self._extract_txt(uploaded)

                if pref_text:
                    st.success("Uploaded file parsed. Prefill suggestions available below.")

            st.markdown("---")

            # Template selector
            template = st.selectbox("Choose a template:", ["Simple", "Modern", "ATS-friendly"], index=0)

            # Personal Info
            st.markdown("### 👤 Personal Information")
            name = st.text_input("Full name", value=self._prefill_field(pref_text, "name") or "")
            email = st.text_input("Email", value=self._prefill_field(pref_text, "email") or "")
            phone = st.text_input("Phone", value=self._prefill_field(pref_text, "phone") or "")
            location = st.text_input("Location (city, country)", value=self._prefill_field(pref_text, "location") or "")

            if email and not self._validate_email(email):
                st.warning("⚠️ The email format looks incorrect.")

            st.markdown("---")

            # Summary
            st.markdown("### 📝 Professional Summary / Objective")
            summary = st.text_area("A single short paragraph (2-4 lines)", value=self._prefill_field(pref_text, "summary") or "", height=90)

            # Skills
            st.markdown("### 🛠️ Skills")
            skills_raw = st.text_input("Comma-separated skills (e.g. Python, ML, SQL)", value=self._prefill_field(pref_text, "skills") or "")
            skills = [s.strip() for s in skills_raw.split(",") if s.strip()]

            st.markdown("---")

            # Experiences (dynamic)
            st.markdown("### 💼 Work Experience (most recent first)")
            self._dynamic_list_ui("cv_experiences", placeholder="Company - Role (YYYY - YYYY) - Key achievements")

            st.markdown("---")

            # Education
            st.markdown("### 🎓 Education")
            self._dynamic_list_ui("cv_educations", placeholder="Degree - Institution (Year) - Notes")

            st.markdown("---")

            # Projects
            st.markdown("### 🚀 Projects / Notable Work")
            self._dynamic_list_ui("cv_projects", placeholder="Project name - short description and tech stack")

        with col_right:
            st.markdown("### Preview")
            # Build preview data dict
            data = {
                "name": name,
                "email": email,
                "phone": phone,
                "location": location,
                "summary": summary,
                "skills": skills,
                "experiences": [e for e in st.session_state.get("cv_experiences", []) if e.strip()],
                "educations": [e for e in st.session_state.get("cv_educations", []) if e.strip()],
                "projects": [p for p in st.session_state.get("cv_projects", []) if p.strip()]
            }

            # Render preview as Markdown (simple formatted view)
            self._render_preview(data, template)

            st.markdown("---")
            st.markdown("### Export")
            colp, cold = st.columns(2)
            with colp:
                if st.button("📄 Download PDF", use_container_width=True):
                    if not name.strip() or not email.strip():
                        st.error("Please provide at least your name and email before exporting.")
                    else:
                        pdf_bytes = self._build_pdf_bytes(data, template)
                        if pdf_bytes:
                            st.download_button(label="⬇️ Download PDF", data=pdf_bytes, file_name=f"{name.replace(' ','_')}_CV.pdf", mime="application/pdf")
                        else:
                            st.error("PDF export is not available (missing dependency 'fpdf'). Install with: pip install fpdf")
            with cold:
                if st.button("📋 Download DOCX", use_container_width=True):
                    if not name.strip() or not email.strip():
                        st.error("Please provide at least your name and email before exporting.")
                    else:
                        docx_bytes = self._build_docx_bytes(data, template)
                        if docx_bytes:
                            st.download_button(label="⬇️ Download DOCX", data=docx_bytes, file_name=f"{name.replace(' ','_')}_CV.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        else:
                            st.error("DOCX export is not available (missing dependency 'python-docx'). Install with: pip install python-docx")

            st.markdown("---")
            st.markdown("### Tips & ATS Score (basic)")
            # Simple ATS friendliness heuristic
            ats_score = self._estimate_ats_score(data)
            st.progress(min(100, max(0, int(ats_score * 100))))
            st.caption("Higher is better. Remove images and complex layouts for ATS-friendly resumes.")

    # ---------------------- Helper UI helpers ----------------------
    def _dynamic_list_ui(self, key: str, placeholder: str = ""):
        items: List[str] = st.session_state.get(key, [""])

        # Display current items with ability to add/remove
        for i in range(len(items)):
            cols = st.columns([5, 1])
            val = st.text_input(f"{key}_{i}", value=items[i], placeholder=placeholder, key=f"{key}_input_{i}")
            items[i] = val
            with cols[1]:
                if st.button("➕", key=f"{key}_add_{i}"):
                    items.insert(i + 1, "")
                    st.session_state[key] = items
                    st.experimental_rerun()
                if st.button("➖", key=f"{key}_remove_{i}"):
                    items.pop(i)
                    st.session_state[key] = items
                    st.experimental_rerun()

        # Ensure at least one empty slot
        if not items:
            items.append("")

        st.session_state[key] = items

    def _render_preview(self, data: dict, template: str = "Simple"):
        # Simple markdown preview; templates can change small style hints
        md = f"# {data.get('name','') if data.get('name') else ''}\n"
        md += f"**{data.get('email','')} | {data.get('phone','')} | {data.get('location','')}**\n\n"

        if data.get('summary'):
            md += f"## Professional Summary\n{data.get('summary')}\n\n"

        if data.get('skills'):
            md += f"**Skills:** {', '.join(data.get('skills'))}\n\n"

        if data.get('experiences'):
            md += "## Work Experience\n"
            for e in data.get('experiences'):
                md += f"- {e}\n"
            md += "\n"

        if data.get('projects'):
            md += "## Projects\n"
            for p in data.get('projects'):
                md += f"- {p}\n"
            md += "\n"

        if data.get('educations'):
            md += "## Education\n"
            for ed in data.get('educations'):
                md += f"- {ed}\n"

        st.markdown(md)

    def _prefill_field(self, text: str, field: str) -> Optional[str]:
        if not text:
            return None
        field = field.lower()
        if field == "email":
            return self._find_email(text)
        if field == "phone":
            return self._find_phone(text)
        if field == "name":
            # Heuristic: first line or largest capitalized sequence
            lines = [l.strip() for l in text.splitlines() if l.strip()]
            if lines:
                # If first line looks like a name (contains letters and spaces, few commas), return it
                first = lines[0]
                if 2 <= len(first.split()) <= 4 and any(c.isalpha() for c in first):
                    return first
            return None
        if field == "summary":
            # Try to extract the first paragraph as summary
            paras = [p.strip() for p in text.split('\n\n') if p.strip()]
            if paras:
                return paras[0][:800]
        if field == "skills":
            # Find lines containing 'Skills' or common skill lists
            m = re.search(r"Skills[:\s\n]+([A-Za-z0-9,\s\-_/+.]+)", text, re.IGNORECASE)
            if m:
                return m.group(1).strip()[:400]
        return None

    def _estimate_ats_score(self, data: dict) -> float:
        # Very basic heuristic: presence of name/email, skills, and plain text sections
        score = 0.0
        if data.get('name'):
            score += 0.2
        if data.get('email') and self._validate_email(data.get('email')):
            score += 0.2
        if data.get('skills'):
            score += 0.2
        if data.get('experiences'):
            score += 0.2
        if data.get('educations'):
            score += 0.1
        return min(1.0, score)
